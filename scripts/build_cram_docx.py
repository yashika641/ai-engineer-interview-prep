"""
Converts Interview-Cram-Sheet.md  ->  Interview-Cram-Sheet.docx

The .md file is the single source of truth (easy to diff / edit in a session).
The .docx is the artefact you actually open on your phone/laptop the day before
an interview. Run this after every edit to the cram sheet.

Supported markdown subset (that's all the cram sheet uses):
  # / ## / ### / ####    headings
  - bullet               (one level of "  - " nesting supported)
  **bold**  and  `code`  inline
  ---                    horizontal rule -> page break
  > quote                callout line
  | tables |             rendered as Word tables

Usage:  python scripts/build_cram_docx.py
Needs:  pip install python-docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "Interview-Cram-Sheet.md"
DOCX = ROOT / "Interview-Cram-Sheet.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(par, text: str) -> None:
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = par.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = par.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            par.add_run(chunk)


def flush_table(doc: Document, buf: list[str]) -> None:
    rows = [r for r in buf if r.strip().startswith("|")]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    for i, r in enumerate(cells):
        wrow = table.add_row().cells
        for j in range(ncol):
            txt = r[j] if j < len(r) else ""
            wrow[j].text = ""
            p = wrow[j].paragraphs[0]
            add_inline(p, txt)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def main() -> None:
    if not MD.exists():
        raise SystemExit(f"Missing {MD}")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    lines = MD.read_text(encoding="utf-8").splitlines()
    tbuf: list[str] = []

    def flush() -> None:
        nonlocal tbuf
        if tbuf:
            flush_table(doc, tbuf)
            tbuf = []

    for raw in lines:
        line = raw.rstrip()

        if line.strip().startswith("|"):
            tbuf.append(line)
            continue
        flush()

        if not line.strip():
            continue
        if line.strip() == "---":
            doc.add_page_break()
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            add_inline(h, m.group(2))
            continue

        if line.lstrip().startswith(("- ", "* ")):
            indent = len(line) - len(line.lstrip())
            style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style_name)
            add_inline(p, line.lstrip()[2:])
            continue

        if line.lstrip().startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run_txt = line.lstrip()[2:]
            add_inline(p, run_txt)
            for run in p.runs:
                run.italic = True
            continue

        p = doc.add_paragraph()
        add_inline(p, line)

    flush()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Interview-Cram-Sheet  -  generated from Interview-Cram-Sheet.md"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX)
    print(f"Wrote {DOCX}")


if __name__ == "__main__":
    main()
